import os
import time
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from PGNDatabase import PGNDatabase


class PGNDocument:
    '''
    Encapsulates a single PGN document
    '''
    def __init__(self, database):
        self.database = database
        self.document = Document()
        self.list_of_games = self.database.get_games()
        self.list_of_drawed_games = self.database.get_draws()


    def create_document(self):
        self.create_document_heading()
        self.create_document_body()
        self.create_document_conclusion()
        self.save_document()

    def create_document_heading(self):
        self.document.add_heading('Chess Database', 0)

    def create_document_introduction(self):
        self.document.add_heading('1. Introduction', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')

    
    def create_document_body(self):
        list_of_games = self.list_of_games
        list_of_drawed_games = self.list_of_drawed_games
        list_of_games_where_stockfish_is_white = self.database.get_games_where_stockfish_is_white()
        list_of_games_where_stockfish_is_black = self.database.get_games_where_stockfish_is_black()
        
        # Todo --> get_stockfish_wins should not return two lists (or should it?)
        list_of_games_where_stockfish_wins_as_white, list_of_games_where_stockfish_wins_as_black = self.database.get_stockfish_wins(list_of_games)
        list_of_games_where_stockfish_losses_as_white, list_of_games_where_stockfish_losses_as_black = self.database.get_stockfish_losses(list_of_games)
        
        list_of_games_where_stockfish_wins = list_of_games_where_stockfish_wins_as_white + list_of_games_where_stockfish_wins_as_black
        list_of_games_where_stockfish_losses = list_of_games_where_stockfish_losses_as_white + list_of_games_where_stockfish_losses_as_black


        self.create_document_introduction()
        # self.create_document_section_for_all_games()
        # Should include a table with game count for each player
        self.document.add_heading('2. Statistics', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.database.get_games())) + ' games. The following sections describe the games in more detail.')
        # self.create_document_subsection_for_all_games()
        

        self.document.add_heading('2.1 General results', level=1)
        self.document.add_paragraph('The following table shows the results of games.')
        self.create_document_result_table_for_all_games(list_of_games, list_of_drawed_games)

        self.document.add_heading('2.1.1 Result table for Stockfish', level=2)
        self.document.add_paragraph('The following table shows the results of games where Stockfish either won or lost, depending on Stockfish color')
        self.document.add_paragraph("def create_document_result_table_with_stockfish(self, list_of_games, list_of_drawed_games):")
        self.create_document_result_table_with_stockfish(list_of_games, list_of_drawed_games)

        self.document.add_heading('2.2 Move count distributions', level=2)
        self.document.add_paragraph('The following graphs shows the distribution of the amount moves in the given set of games.')
                
        self.document.add_heading('2.2.1 All games', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in all games.')
        # Plot Cumulative Moves Distribution for all games, games where Stockfish is white and games where Stockfish is black
        dictionary_for_first_plot = {"All games": list_of_games, "Games where Stockfish is white": list_of_games_where_stockfish_is_white, "Games where Stockfish is black": list_of_games_where_stockfish_is_black}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_first_plot, "1stMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for all games')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games)

        self.document.add_heading('2.2.2 Either stockfish won or draws', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish won.')
        # Plot Cumulative Moves Distribution for games where Stockfish won and games where Stockfish lost
        dictionary_for_second_plot = {"Games where Stockfish won": list_of_games_where_stockfish_wins, "Games that ended in draw": list_of_drawed_games}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_second_plot, "2ndMoveCountCDPlot.png")
        
        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish won or drew')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games_where_stockfish_wins)
        self.document.add_paragraph("A noteworthy observation is that we get a spike in the distribution of moves when Stockfish draws.")
                                    
        self.document.add_heading('2.2.3 Stockfish loses', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish lost.')
        dictionary_for_third_plot = {"Games where Stockfish lost": list_of_games_where_stockfish_losses}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_third_plot, "3rdMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish lost')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games_where_stockfish_losses)
        
        self.document.add_heading('2.3 Plycount distribution', level=2)
        self.document.add_paragraph('The following graph shows the distribution of plycount in the database.')
        self.document.add_paragraph('The x-axis shows the plycount, and the y-axis shows the number of games with that plycount.')
        self.add_picture_of_plycount_distribution()



    def create_document_section_for_all_games(self):
        # Should include a table with game count for each player
        self.document.add_heading('2 Games', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.database.get_games())) + ' games. The following sections describe the games in more detail.')

    def create_document_subsection_for_all_games(self):
        return 

    def create_document_result_table_for_all_games(self, games, draws):
        
        black_wins = self.database.get_black_wins()
        white_wins = self.database.get_white_wins()
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'White wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Black wins'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(white_wins))
        row_cells[2].text = str(len(draws))
        row_cells[3].text = str(len(black_wins))
        row_cells = table.add_row().cells
        row_cells[0].text = 'Percentage'
        row_cells[1].text = str(round((len(white_wins) / len(games)) * 100, 2)) + '%'
        row_cells[2].text = str(round((len(draws) / len(games)) * 100, 2)) + '%'
        row_cells[3].text = str(round((len(black_wins) / len(games)) * 100, 2)) + '%'


    def create_document_result_table_with_stockfish(self, list_of_games, list_of_drawed_games):
        stockfish_wins_as_white, stockfish_wins_as_black = self.database.get_stockfish_wins(list_of_games)
        stockfish_losses_as_white, stockfish_losses_as_black = self.database.get_stockfish_losses(list_of_games)
        stockfish_draws_as_white, stockfish_draws_as_black = self.database.get_stockfish_draws(list_of_drawed_games)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Losses'
        hdr_cells[4].text = 'Winning Percentage'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as white'
        row_cells[1].text = str(len(stockfish_wins_as_white))
        row_cells[2].text = str(len(stockfish_draws_as_white))
        row_cells[3].text = str(len(stockfish_losses_as_white))
        row_cells[4].text = str(round((len(stockfish_wins_as_white) / (len(stockfish_wins_as_white) + len(stockfish_draws_as_white) + len(stockfish_losses_as_white))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as black'
        row_cells[1].text = str(len(stockfish_wins_as_black))
        row_cells[2].text = str(len(stockfish_draws_as_black))
        row_cells[3].text = str(len(stockfish_losses_as_black))
        row_cells[4].text = str(round((len(stockfish_wins_as_black) / (len(stockfish_wins_as_black) + len(stockfish_draws_as_black) + len(stockfish_losses_as_black))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(stockfish_wins_as_white) + len(stockfish_wins_as_black))
        row_cells[2].text = str(len(stockfish_draws_as_white) + len(stockfish_draws_as_black))
        row_cells[3].text = str(len(stockfish_losses_as_white) + len(stockfish_losses_as_black))
        row_cells[4].text = str(round(((len(stockfish_wins_as_white) + len(stockfish_wins_as_black)) / len(list_of_games)) * 100, 2)) + '%'


    def add_picture_of_plycount_distribution(self):
        list_of_games = self.database.get_games()
        self.database.plot_plycount_distribution(list_of_games)
        self.document.add_picture('plycount_distribution.png', width=Inches(6))


    def create_document_moves_distribution(self, list_of_games, list_of_games_where_stockfish_is_white, list_of_games_where_stockfish_is_black):
        self.document.add_heading('2.4 Moves distribution', level=2)
        self.document.add_paragraph('The following graph shows the distribution of moves in the database.')
        self.document.add_paragraph('The x-axis shows the number of moves, and the y-axis shows the number of games with that number of moves.')
        self.database.plot_move_count_distribution(list_of_games_where_stockfish_is_white, "Stockfish as white")
        self.database.plot_move_count_distribution(list_of_games_where_stockfish_is_black, "Stockfish as black")
        self.database.plot_move_count_distribution(list_of_games, "All games")
        self.document.add_picture('move_count_distribution.png', width=Inches(6))


    def add_picture_of_cumulative_moves_distribution_for_multiple_games(self, dict, filename): # dict is a dictionary with the key being the name of the list of games, and the value being the list of games
        fig, ax = plt.subplots()
        fig.set_size_inches(10,5)
        # Iterate through the dictionary, and plot each list of games
        self.database.plot_multiple_move_count_histogram_cumulative(dict)
        plt.savefig(filename)
        self.document.add_picture(filename, width=Inches(6))
        


    ### This is the old version of the function above, which only plots one list of games at a time ###
    def create_document_cumulative_moves_distribution(self, list_of_games, list_of_games_where_stockfish_is_white, list_of_games_where_stockfish_is_black):
        # self.database.clear_plot()
        self.document.add_heading('2.4 Moves distribution', level=2)
        self.document.add_paragraph('The following graph shows the distribution of moves in the database.')
        self.document.add_paragraph('The x-axis shows the number of moves, and the y-axis shows the number of games with that number of moves.')
        fig, ax = plt.subplots()
        fig.set_size_inches(10,5)
        
        self.database.plot_move_count_histogram_cumulative(list_of_games_where_stockfish_is_white, "Stockfish as white", axis=ax)
        self.database.plot_move_count_histogram_cumulative(list_of_games_where_stockfish_is_black, "Stockfish as black", axis=ax)
        self.database.plot_move_count_histogram_cumulative(list_of_games, "All games", axis=ax)
        plt.legend()
        self.database.save_histogram('move_count_distribution.png')
        self.document.add_picture('move_count_distribution.png', width=Inches(6))

    def create_document_moves_table_mean_and_standard_deviation(self):
        self.document.add_heading('2.4.1 Moves table', level=2)
        self.document.add_paragraph('The following table shows the mean and standard deviation of the number of moves in the database.')
        

    def add_table_of_mean_and_standard_deviation_of_moves(self, list_of_games):
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Standard Deviation'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Number of moves'
        row_cells[1].text = str(round(self.database.get_mean_number_of_moves(list_of_games), 2))
        row_cells[2].text = str(round(self.database.get_standard_deviation_of_moves(list_of_games), 2))
        

    def create_document_conclusion(self):
        self.document.add_heading('3. Conclusion', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')


    def save_document(self):
        if os.path.exists('ChessDatabase.docx'):
            os.remove('ChessDatabase.docx') # Remove the file if it already exists to avoid an error
        self.document.save('ChessDatabase.docx')

def main():
    start_time = time.time()
    database = PGNDatabase("./big_database.pgn")
    document = PGNDocument(database)
    print("Time elapsed: " + (str(time.time() - start_time)) + " seconds")
    document.create_document()
    print("Time elapsed: " + (str(time.time() - start_time)) + " seconds")



if __name__ == "__main__":
    main()




    
