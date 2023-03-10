import os
import time
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from Database import Database
from Tree import *
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from Plot import *

class PGNDocument:
    '''
    Encapsulates a single PGN document
    '''
    def __init__(self, database, opening_occurrences, include_openings, max_tree_depth, minimum_games_on_node_to_keep_going_on_a_branch):
        
        
        self.database = database
        self.document = Document()
        self.opening_occurrences = opening_occurrences
        self.include_openings = include_openings
        self.max_tree_depth = max_tree_depth
        self.minimum_games_on_node_to_keep_going_on_a_branch = minimum_games_on_node_to_keep_going_on_a_branch
        self.list_of_games = self.database.get_list_of_games()

        # TODO: skriv i dokumentet hvorfor vi iterer over games 12,13 ganger mer enn vi trenger

        self.list_of_games = self.database.get_list_of_games()
        self.list_of_drawed_games = self.database.get_list_of_draws()
        self.list_of_games_where_stockfish_is_white = self.database.get_list_of_games_where_stockfish_is_white()
        self.list_of_games_where_stockfish_is_black = self.database.get_list_of_games_where_stockfish_is_black()
        self.list_of_games_where_stockfish_wins = self.database.get_list_of_stockfish_wins()
        self.list_of_games_where_stockfish_losses = self.database.get_list_of_stockfish_losses()  
        self.list_of_games_where_stockfish_wins_or_draws = self.database.get_list_of_games_where_stockfish_wins_or_draws()

        self.stockfish_wins_as_white = self.database.get_list_of_stockfish_wins_as_white()
        self.stockfish_wins_as_black = self.database.get_list_of_stockfish_wins_as_black()
        self.stockfish_losses_as_white = self.database.get_list_of_stockfish_losses_as_white()
        self.stockfish_losses_as_black = self.database.get_list_of_stockfish_losses_as_black()
        self.stockfish_draws_as_white = self.database.get_list_of_stockfish_draws_as_white()
        self.stockfish_draws_as_black = self.database.get_list_of_stockfish_draws_as_black()
        

        self.black_wins = self.database.get_list_of_black_wins()
        self.white_wins = self.database.get_list_of_white_wins()
        self.draws = self.database.get_list_of_draws()

        self.database_of_games_where_stockfish_wins_or_draws = Database(self.list_of_games_where_stockfish_wins_or_draws)
        self.database_of_games_where_stockfish_wins = Database(self.list_of_games_where_stockfish_wins)
        self.database_of_games_where_stockfish_losses = Database(self.list_of_games_where_stockfish_losses)

    def create_document(self):
        self.create_document_heading()
        self.create_document_body()
        self.create_document_conclusion()
        self.save_document()


    ## MAIN COMPONENTS ##

    def create_document_heading(self):
        self.document.add_heading('Chess Report', 0)

    def create_document_body(self):
        self.create_document_introduction()

        self.document.add_heading('2. Statistics', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.list_of_games)) + ' games. The following sections describe the games in more detail.')

        self.document.add_heading('2.1 General results', level=1)
        self.document.add_paragraph('The following table shows the results of games.')
        self.create_document_result_table_for_all_games()

        self.document.add_heading('2.1.1 Result table for Stockfish', level=2)
        self.document.add_paragraph('The following table shows the results of games where Stockfish either won or lost, depending on Stockfish color')
        self.create_document_result_table_with_stockfish()

        self.document.add_heading('2.2 Move count distributions', level=2)
        self.document.add_paragraph('The following graphs shows the distribution of the amount moves in the given set of games.')
                
        self.document.add_heading('2.2.1 All games', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in all games.')



        # Plot Cumulative Moves Distribution for all games, games where Stockfish is white and games where Stockfish is black
        dictionary_for_first_plot = {"All games": self.list_of_games, "Games where Stockfish is white": self.list_of_games_where_stockfish_is_white, "Games where Stockfish is black": self.list_of_games_where_stockfish_is_black}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_first_plot, "./plots/1stMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for all games')
        self.add_table_of_mean_and_standard_deviation_of_moves(self.database)

        self.document.add_heading('2.2.2 Either stockfish won or draws', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish won.')
        # Plot Cumulative Moves Distribution for games where Stockfish won and games where Stockfish lost
        dictionary_for_second_plot = {"Games where Stockfish won": self.list_of_games_where_stockfish_wins, "Games that ended in draw": self.list_of_drawed_games}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_second_plot, "./plots/2ndMoveCountCDPlot.png")
        

        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish won or drew')
        self.add_table_of_mean_and_standard_deviation_of_moves(self.database_of_games_where_stockfish_wins_or_draws)
        self.document.add_paragraph("A noteworthy observation is that we get a spike in the distribution of moves when Stockfish draws.")
                                    
        self.document.add_heading('2.2.3 Stockfish loses', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish lost.')
        dictionary_for_third_plot = {"Games where Stockfish lost": self.list_of_games_where_stockfish_losses}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_third_plot, "./plots/3rdMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish lost')
        self.add_table_of_mean_and_standard_deviation_of_moves(self.database_of_games_where_stockfish_losses)

        self.create_openings_table()

        self.create_document_section_for_tree_plotting()
    
    def create_document_conclusion(self):
        self.document.add_heading('3. Conclusion', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')

    def save_document(self):
        if os.path.exists('ChessReport.docx'):
            os.remove('ChessReport.docx') # Remove the file if it already exists to avoid an error
        self.document.save('ChessReport.docx')


    ## SUBCOMPONENTS ##
    def create_document_introduction(self):
        self.document.add_heading('1. Introduction', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')

    def create_openings_table(self):
        self.document.add_heading('3.1 Openings', level=2)
        self.document.add_paragraph('The following table shows the openings that occured at least ' + str(self.opening_occurrences) + ' times.')
        openings = self.database.get_openings_that_occurred_at_least_n_times(self.opening_occurrences)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Opening'
        hdr_cells[1].text = 'White wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Black wins'
        hdr_cells[4].text = 'Total games'
        for opening in openings:
            opening_list = self.database.get_list_with_opening(opening)
            opening_database = Database(opening_list)
            white_wins = opening_database.get_list_of_white_wins()
            black_wins = opening_database.get_list_of_black_wins()
            draws = opening_database.get_list_of_draws()
            list_of_games = opening_database.get_list_of_games()
            row_cells = table.add_row().cells
            row_cells[0].text = opening
            row_cells[1].text = str(len(white_wins))
            row_cells[2].text = str(len(draws))
            row_cells[3].text = str(len(black_wins))
            row_cells[4].text = str(len(list_of_games))

    def create_document_section_for_tree_plotting(self):
        self.document.add_page_break()
        self.document.add_heading('3 Tree plotting', level=1)
        self.document.add_paragraph('The following section describes the tree plotting. The tree plotting is done using the Tree class........')
        self.document.add_paragraph('We choose to plot the following trees with depth 10, first the Sicilian defence, then the French defence.')
        
        # TODO: legg til funksjonalitet at man

        openings = self.database.get_openings_that_occurred_at_least_n_times(self.opening_occurrences)
        self.document.add_page_break()

        for opening in openings:
            if opening not in self.include_openings:
                continue

            self.document.add_heading(opening, level=2)
            list_of_games = self.database.get_list_with_opening(opening)
            opening_filename = opening.lower().replace(" ", "_").replace("'", "")
            
            tree = OpeningTree(list_of_games)
            tree.save_tree(self.max_tree_depth, self.minimum_games_on_node_to_keep_going_on_a_branch ,opening_filename)

            p = self.document.add_paragraph('')
            print("./graphs/ " + opening_filename + ".png")
            self.add_hyperlink(p, 'open full picture', "./graphs/" + opening_filename + ".png")
            self.document.add_picture("./graphs/" + opening_filename + ".png", width=Inches(6))
            self.document.add_page_break()

    def create_document_section_for_all_games(self):
        # Should include a table with game count for each player
        self.document.add_heading('2 Games', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.database.get_games())) + ' games. The following sections describe the games in more detail.')

    def create_document_subsection_for_all_games(self):
        return 

    def create_document_result_table_for_all_games(self):
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'White wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Black wins'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(self.white_wins))
        row_cells[2].text = str(len(self.draws))
        row_cells[3].text = str(len(self.black_wins))
        row_cells = table.add_row().cells
        row_cells[0].text = 'Percentage'
        row_cells[1].text = str(self.database.get_precentage_of_white_wins()) + '%'
        row_cells[2].text = str(self.database.get_precentage_of_draws()) + '%'
        row_cells[3].text = str(self.database.get_precentage_of_black_wins()) + '%'

    def create_document_result_table_with_stockfish(self):
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Losses'
        hdr_cells[4].text = 'Winning Percentage'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as white'
        row_cells[1].text = str(len(self.stockfish_wins_as_white))
        row_cells[2].text = str(len(self.stockfish_draws_as_white))
        row_cells[3].text = str(len(self.stockfish_losses_as_white))
        row_cells[4].text = str(round((len(self.stockfish_wins_as_white) / (len(self.stockfish_wins_as_white) + len(self.stockfish_draws_as_white) + len(self.stockfish_losses_as_white))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as black'
        row_cells[1].text = str(len(self.stockfish_wins_as_black))
        row_cells[2].text = str(len(self.stockfish_draws_as_black))
        row_cells[3].text = str(len(self.stockfish_losses_as_black))
        row_cells[4].text = str(round((len(self.stockfish_wins_as_black) / (len(self.stockfish_wins_as_black) + len(self.stockfish_draws_as_black) + len(self.stockfish_losses_as_black))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(self.stockfish_wins_as_white) + len(self.stockfish_wins_as_black))
        row_cells[2].text = str(len(self.stockfish_draws_as_white) + len(self.stockfish_draws_as_black))
        row_cells[3].text = str(len(self.stockfish_losses_as_white) + len(self.stockfish_losses_as_black))
        row_cells[4].text = str(round(((len(self.stockfish_wins_as_white) + len(self.stockfish_wins_as_black)) / len(self.list_of_games)) * 100, 2)) + '%'

    def add_picture_of_cumulative_moves_distribution_for_multiple_games(self, dict, filename): # dict is a dictionary with the key being the name of the list of games, and the value being the list of games
        plot = Plot()
        plot.plot_multiple_move_count_histogram_cumulative(dict, filename)
        self.document.add_picture(filename, width=Inches(6))

    def create_document_moves_table_mean_and_standard_deviation(self):
        self.document.add_heading('2.4.1 Moves table', level=2)
        self.document.add_paragraph('The following table shows the mean and standard deviation of the number of moves in the database.')
        
    
    def add_table_of_mean_and_standard_deviation_of_moves(self, database):
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Standard Deviation'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Number of moves'
        row_cells[1].text = str(round(database.get_mean_number_of_moves(), 2))
        row_cells[2].text = str(round(database.get_standard_deviation_of_moves(), 2))
        

    ## HELPER FUNCTIONS ##
    
    # TODO, vise kilde til denne?

    def add_hyperlink(self, paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink


def main():
    start_time = time.time()

    database = Database()
    database.parse_from_pgn("./databases/Stockfish_15_64-bit.commented.[2600].pgn")
    
    # parametere til Document 
    minimum_opening_occurences_to_add_to_table = 30
    max_tree_depth = 15
    minimum_games_on_node_to_keep_going_on_a_branch = 4

    inlcude_opening_graphs = ["Nimzo-Indian", "Sicilian", "Sicilian defence" ,"Ruy Lopez", "King's Indian"]

    document = PGNDocument(database, minimum_opening_occurences_to_add_to_table, inlcude_opening_graphs, max_tree_depth, minimum_games_on_node_to_keep_going_on_a_branch)
    document.create_document()

if __name__ == "__main__":
    main()




    
