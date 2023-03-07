import os
import time
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from PGNDatabase import PGNDatabase
from Tree import *
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

class PGNDocument:
    '''
    Encapsulates a single PGN document
    '''
    def __init__(self, database, opening_occurrences):
        self.database = database
        self.document = Document()
        self.opening_occurrences = opening_occurrences
        self.list_of_games = self.database.get_games()
        self.list_of_drawed_games = self.database.get_draws(self.list_of_games)


    def create_document(self):
        self.create_document_heading()
        self.create_document_body()
        self.create_document_conclusion()
        self.save_document()

    def create_document_heading(self):
        self.document.add_heading('Chess Database', 0)

    def create_document_introduction(self):
        self.document.add_heading('1. Introduction', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')

    
    def create_document_body(self):
        list_of_games = self.list_of_games
        list_of_drawed_games = self.list_of_drawed_games
        list_of_games_where_stockfish_is_white = self.database.get_games_where_stockfish_is_white()
        list_of_games_where_stockfish_is_black = self.database.get_games_where_stockfish_is_black()
        
        list_of_games_where_stockfish_wins_as_white, list_of_games_where_stockfish_wins_as_black = self.database.get_stockfish_wins(list_of_games)
        list_of_games_where_stockfish_losses_as_white, list_of_games_where_stockfish_losses_as_black = self.database.get_stockfish_losses(list_of_games)
        
        list_of_games_where_stockfish_wins = list_of_games_where_stockfish_wins_as_white + list_of_games_where_stockfish_wins_as_black
        list_of_games_where_stockfish_losses = list_of_games_where_stockfish_losses_as_white + list_of_games_where_stockfish_losses_as_black


        self.create_document_introduction()
        # self.create_document_section_for_all_games()
        # Should include a table with game count for each player
        self.document.add_heading('2. Statistics', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.database.get_games())) + ' games. The following sections describe the games in more detail.')
        # self.create_document_subsection_for_all_games()

        self.document.add_heading('2.1 General results', level=1)
        self.document.add_paragraph('The following table shows the results of games.')
        self.create_document_result_table_for_all_games(list_of_games, list_of_drawed_games)

        self.document.add_heading('2.1.1 Result table for Stockfish', level=2)
        self.document.add_paragraph('The following table shows the results of games where Stockfish either won or lost, depending on Stockfish color')
        self.document.add_paragraph("def create_document_result_table_with_stockfish(self, list_of_games, list_of_drawed_games):")
        self.create_document_result_table_with_stockfish(list_of_games, list_of_drawed_games)

        self.document.add_heading('2.2 Move count distributions', level=2)
        self.document.add_paragraph('The following graphs shows the distribution of the amount moves in the given set of games.')
                
        self.document.add_heading('2.2.1 All games', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in all games.')
        # Plot Cumulative Moves Distribution for all games, games where Stockfish is white and games where Stockfish is black
        dictionary_for_first_plot = {"All games": list_of_games, "Games where Stockfish is white": list_of_games_where_stockfish_is_white, "Games where Stockfish is black": list_of_games_where_stockfish_is_black}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_first_plot, "./plots/1stMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for all games')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games)

        self.document.add_heading('2.2.2 Either stockfish won or draws', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish won.')
        # Plot Cumulative Moves Distribution for games where Stockfish won and games where Stockfish lost
        dictionary_for_second_plot = {"Games where Stockfish won": list_of_games_where_stockfish_wins, "Games that ended in draw": list_of_drawed_games}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_second_plot, "./plots/2ndMoveCountCDPlot.png")
        
        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish won or drew')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games_where_stockfish_wins)
        self.document.add_paragraph("A noteworthy observation is that we get a spike in the distribution of moves when Stockfish draws.")
                                    
        self.document.add_heading('2.2.3 Stockfish loses', level=3)
        self.document.add_paragraph('The following graph shows the distribution of the amount moves in games where Stockfish lost.')
        dictionary_for_third_plot = {"Games where Stockfish lost": list_of_games_where_stockfish_losses}
        self.add_picture_of_cumulative_moves_distribution_for_multiple_games(dictionary_for_third_plot, "./plots/3rdMoveCountCDPlot.png")
        self.document.add_paragraph('Mean and standard deviation table for games where Stockfish lost')
        self.add_table_of_mean_and_standard_deviation_of_moves(list_of_games_where_stockfish_losses)


        ################## Openings table. Opening name, number of games, white wins, black wins, draws ############################
        self.document.add_heading('3.1 Openings', level=2)
        self.document.add_paragraph('The following table shows the openings that occured at least ' + str(self.opening_occurrences) + ' times.')
        openings = self.database.get_openings_that_occurred_at_least_n_times(self.opening_occurrences)
        print(openings)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Opening'
        hdr_cells[1].text = 'White wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Black wins'
        hdr_cells[4].text = 'Total games'
        for opening in openings:
            list_of_games = self.database.get_games_with_opening(opening)
            white_wins = self.database.get_white_wins(list_of_games)
            black_wins = self.database.get_black_wins(list_of_games)
            draws = self.database.get_draws(list_of_games)
            row_cells = table.add_row().cells
            row_cells[0].text = opening
            row_cells[1].text = str(len(white_wins))
            row_cells[2].text = str(len(draws))
            row_cells[3].text = str(len(black_wins))
            row_cells[4].text = str(len(list_of_games))

        self.create_document_section_for_tree_plotting()

    def add_hyperlink(self, paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element and a new w:rPr element
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        # Create a new Run object and add the hyperlink into it
        r = paragraph.add_run ()
        r._r.append (hyperlink)

        # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
        # Delete this if using a template that has the hyperlink style in it
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True

        return hyperlink


    ############ TREE PLOTTING ############################
    def create_document_section_for_tree_plotting(self):
        self.document.add_heading('3 Tree plotting', level=1)
        self.document.add_paragraph('The following section describes the tree plotting. The tree plotting is done using the Tree class........')
        self.document.add_paragraph('We choose to plot the following trees with depth 10, first the Sicilian defence, then the French defence.')
        
        self.document.add_heading('3.1 Sicilian defence', level=2)
        list_of_games = self.database.get_games_with_opening("Sicilian defence")
        save_tree_from_list_of_games(list_of_games, 10, "tree_Sicilian")
        p = self.document.add_paragraph('')
        self.add_hyperlink(p, 'open full picture', "./graphs/tree_Sicilian.png")
        self.document.add_picture("./graphs/tree_Sicilian.png", width=Inches(8.0))

        self.document.add_heading('3.2 French defence', level=2)
        list_of_games = self.database.get_games_with_opening("French")
        save_tree_from_list_of_games(list_of_games, 10, "tree_French")
        self.document.add_picture("./graphs/tree_French.png", width=Inches(8.0))

        self.document.add_heading('3.3 Birds opening', level=2)
        list_of_games = self.database.get_games_with_opening("Bird's opening")
        save_tree_from_list_of_games(list_of_games, 3, "tree_Bird")
        self.document.add_picture("./graphs/tree_Bird.png", width=Inches(8.0))







        
        # self.document.add_heading('2.3 Plycount distribution', level=2)
        # self.document.add_paragraph('The following graph shows the distribution of plycount in the database.')
        # self.document.add_paragraph('The x-axis shows the plycount, and the y-axis shows the number of games with that plycount.')
        # self.add_picture_of_plycount_distribution()



    def create_document_section_for_all_games(self):
        # Should include a table with game count for each player
        self.document.add_heading('2 Games', level=2)
        self.document.add_paragraph('The database contains ' + str(len(self.database.get_games())) + ' games. The following sections describe the games in more detail.')

    def create_document_subsection_for_all_games(self):
        return 

    def create_document_result_table_for_all_games(self, games, draws):
        black_wins = self.database.get_black_wins(games)
        white_wins = self.database.get_white_wins(games)
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'White wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Black wins'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(white_wins))
        row_cells[2].text = str(len(draws))
        row_cells[3].text = str(len(black_wins))
        row_cells = table.add_row().cells
        row_cells[0].text = 'Percentage'
        row_cells[1].text = str(round((len(white_wins) / len(games)) * 100, 2)) + '%'
        row_cells[2].text = str(round((len(draws) / len(games)) * 100, 2)) + '%'
        row_cells[3].text = str(round((len(black_wins) / len(games)) * 100, 2)) + '%'


    def create_document_result_table_with_stockfish(self, list_of_games, list_of_drawed_games):
        stockfish_wins_as_white, stockfish_wins_as_black = self.database.get_stockfish_wins(list_of_games)
        stockfish_losses_as_white, stockfish_losses_as_black = self.database.get_stockfish_losses(list_of_games)
        stockfish_draws_as_white, stockfish_draws_as_black = self.database.get_stockfish_draws(list_of_drawed_games)
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Wins'
        hdr_cells[2].text = 'Draws'
        hdr_cells[3].text = 'Losses'
        hdr_cells[4].text = 'Winning Percentage'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as white'
        row_cells[1].text = str(len(stockfish_wins_as_white))
        row_cells[2].text = str(len(stockfish_draws_as_white))
        row_cells[3].text = str(len(stockfish_losses_as_white))
        row_cells[4].text = str(round((len(stockfish_wins_as_white) / (len(stockfish_wins_as_white) + len(stockfish_draws_as_white) + len(stockfish_losses_as_white))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Starts as black'
        row_cells[1].text = str(len(stockfish_wins_as_black))
        row_cells[2].text = str(len(stockfish_draws_as_black))
        row_cells[3].text = str(len(stockfish_losses_as_black))
        row_cells[4].text = str(round((len(stockfish_wins_as_black) / (len(stockfish_wins_as_black) + len(stockfish_draws_as_black) + len(stockfish_losses_as_black))) * 100, 2)) + '%'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Total'
        row_cells[1].text = str(len(stockfish_wins_as_white) + len(stockfish_wins_as_black))
        row_cells[2].text = str(len(stockfish_draws_as_white) + len(stockfish_draws_as_black))
        row_cells[3].text = str(len(stockfish_losses_as_white) + len(stockfish_losses_as_black))
        row_cells[4].text = str(round(((len(stockfish_wins_as_white) + len(stockfish_wins_as_black)) / len(list_of_games)) * 100, 2)) + '%'


    def add_picture_of_plycount_distribution(self):
        list_of_games = self.database.get_games()
        self.database.plot_plycount_distribution(list_of_games)
        self.document.add_picture('./plots/plycount_distribution.png', width=Inches(6))


    def create_document_moves_distribution(self, list_of_games, list_of_games_where_stockfish_is_white, list_of_games_where_stockfish_is_black):
        self.document.add_heading('2.4 Moves distribution', level=2)
        self.document.add_paragraph('The following graph shows the distribution of moves in the database.')
        self.document.add_paragraph('The x-axis shows the number of moves, and the y-axis shows the number of games with that number of moves.')
        self.database.plot_move_count_distribution(list_of_games_where_stockfish_is_white, "Stockfish as white")
        self.database.plot_move_count_distribution(list_of_games_where_stockfish_is_black, "Stockfish as black")
        self.database.plot_move_count_distribution(list_of_games, "All games")
        self.document.add_picture('./plots/move_count_distribution.png', width=Inches(6))


    def add_picture_of_cumulative_moves_distribution_for_multiple_games(self, dict, filename): # dict is a dictionary with the key being the name of the list of games, and the value being the list of games
        fig, ax = plt.subplots()
        fig.set_size_inches(10,5)
        # Iterate through the dictionary, and plot each list of games
        self.database.plot_multiple_move_count_histogram_cumulative(dict)
        plt.savefig(filename)
        self.document.add_picture(filename, width=Inches(6))
        


    ### This is the old version of the function above, which only plots one list of games at a time ###
    def create_document_cumulative_moves_distribution(self, list_of_games, list_of_games_where_stockfish_is_white, list_of_games_where_stockfish_is_black):
        # self.database.clear_plot()
        self.document.add_heading('2.4 Moves distribution', level=2)
        self.document.add_paragraph('The following graph shows the distribution of moves in the database.')
        self.document.add_paragraph('The x-axis shows the number of moves, and the y-axis shows the number of games with that number of moves.')
        fig, ax = plt.subplots()
        fig.set_size_inches(10,5)
        
        self.database.plot_move_count_histogram_cumulative(list_of_games_where_stockfish_is_white, "Stockfish as white", axis=ax)
        self.database.plot_move_count_histogram_cumulative(list_of_games_where_stockfish_is_black, "Stockfish as black", axis=ax)
        self.database.plot_move_count_histogram_cumulative(list_of_games, "All games", axis=ax)
        plt.legend()
        self.database.save_histogram('move_count_distribution.png')
        self.document.add_picture('move_count_distribution.png', width=Inches(6))

    def create_document_moves_table_mean_and_standard_deviation(self):
        self.document.add_heading('2.4.1 Moves table', level=2)
        self.document.add_paragraph('The following table shows the mean and standard deviation of the number of moves in the database.')
        

    def add_table_of_mean_and_standard_deviation_of_moves(self, list_of_games):
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''
        hdr_cells[1].text = 'Mean'
        hdr_cells[2].text = 'Standard Deviation'
        row_cells = table.add_row().cells
        row_cells[0].text = 'Number of moves'
        row_cells[1].text = str(round(self.database.get_mean_number_of_moves(list_of_games), 2))
        row_cells[2].text = str(round(self.database.get_standard_deviation_of_moves(list_of_games), 2))
        

    def create_document_conclusion(self):
        self.document.add_heading('3. Conclusion', level=1)
        self.document.add_paragraph('This document is a summary of the chess database.')


    def save_document(self):
        if os.path.exists('ChessDatabase.docx'):
            os.remove('ChessDatabase.docx') # Remove the file if it already exists to avoid an error
        self.document.save('ChessDatabase.docx')

def main():
    start_time = time.time()
    database = PGNDatabase("./databases/Stockfish_15_64-bit.commented.[2600].pgn")
    opening_occurrences = 40
    document = PGNDocument(database, opening_occurrences)



    print("Time elapsed: " + (str(time.time() - start_time)) + " seconds")
    document.create_document()
    print("Time elapsed: " + (str(time.time() - start_time)) + " seconds")



if __name__ == "__main__":
    main()




    
